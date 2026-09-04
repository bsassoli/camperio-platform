# -*- coding: utf-8 -*-
"""Motore look-through valutario Camperio (metodologia skill camperio-portfolio).
Calcola la distribuzione FX dei 3 fondi UCITS dai file Controlfida e costruisce la
matrice valutaria del portafoglio (azioni dirette + ETF + fondi look-through + bond + liquidità).
Oro = asset class autonoma (fuori FX). Derivati overlay = informativi (non FX)."""
import os, numbers
from collections import defaultdict
from camperio_core.portfolio import methodology as M
HERE = os.path.dirname(os.path.abspath(__file__))

# colonne (0-based) del tracciato Controlfida (stabile)
C_DIV, C_ESP, C_TIPS, C_PAESE, C_MSEC, C_DESC = 9, 22, 26, 50, 63, 2

# sigla_paese -> valuta (per Superdiscovery)
CMAP = {'CN':'CNY','HK':'HKD','TW':'TWD','KR':'KRW','IN':'INR','BR':'BRL','MX':'MXN','ID':'IDR',
        'ZA':'ZAR','TR':'TRY','PL':'PLN','HU':'HUF','CZ':'CZK','TH':'THB','MY':'MYR','PH':'PHP',
        'QA':'QAR','AE':'AED','SA':'SAR','EG':'EGP','GR':'EUR','CL':'CLP','CO':'COP','PE':'PEN',
        'US':'USD','GB':'GBP'}

# codabi fondi nel portafoglio Camperio -> tipo fondo
FUND_CODE = {"95ZYC2": "DELTA", "8W8C01": "SUPERDISCOVERY", "54OPLF": "ALPHAGREEN", "8W8C34": "DELTADEF"}

def _rows(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    ws = wb[wb.sheetnames[0]]
    return list(ws.iter_rows(min_row=6, values_only=True))

def _num(v):
    return v if isinstance(v, numbers.Number) else 0

def fund_fx_fractions(path, tipo):
    """Ritorna {valuta: frazione} (somma=1) per un fondo, secondo la regola asimmetrica.
    tipo in {DELTA, DELTADEF, SUPERDISCOVERY, ALPHAGREEN}."""
    rows = _rows(path)
    fx = defaultdict(float)
    if tipo in ("DELTA", "DELTADEF"):
        phys = {"Bond Governative", "Bond", "Cash", "Formation Expenses", "Variation Margins", "Committed Liquidity"}
        eurusd = chf = 0.0
        for r in rows:
            t = r[C_TIPS]; e = _num(r[C_ESP])
            if not t: continue
            div = (r[C_DIV] or "").strip()
            if t in phys:
                fx[div or "EUR"] += e
            elif t in ("Option", "Future"):
                msec = r[C_MSEC] or ""
                if msec == "Curncy":
                    if "CHF" in str(r[C_DESC]).upper(): chf += e
                    else: eurusd += e
                # Index (overlay azionario) e Comdty (Bobl) -> esclusi
        fx["USD"] -= eurusd; fx["EUR"] += eurusd
        fx["USD"] -= chf;    fx["CHF"] += chf
    elif tipo == "SUPERDISCOVERY":
        for r in rows:
            t = r[C_TIPS]; e = _num(r[C_ESP])
            if not e: continue
            div = (r[C_DIV] or "").strip(); paese = (r[C_PAESE] or "").strip().upper()
            if t in ("Equity", "ETF"):
                ccy = CMAP.get(paese) or (div if div else "EM")
                if ccy == "USD": ccy = "EM"
                fx["Other EM" if ccy == "EM" else ccy] += e
            elif t == "Future":
                continue  # MXEF overlay escluso dalla FX
            elif t in ("Bond Governative", "Bond", "Cash", "Variation Margins", "Committed Liquidity", "Formation Expenses"):
                fx["Other EM" if (div == "USD" or div == "") else div] += e
    elif tipo == "ALPHAGREEN":
        for r in rows:
            t = r[C_TIPS]; e = _num(r[C_ESP])
            if not e or t is None: continue
            fx[(r[C_DIV] or "EUR").strip()] += e
    tot = sum(fx.values()) or 1.0
    return {k: v / tot for k, v in fx.items() if abs(v) > 1e-9}

def _find_fund_file(repo_dir, prefix):
    if not os.path.isdir(repo_dir): return None
    cand = sorted([f for f in os.listdir(repo_dir)
                   if f.upper().replace("_", "").startswith(prefix) and f.lower().endswith((".xlsx", ".xls"))
                   and not (prefix == "DELTA" and f.upper().replace("_", "").startswith("DELTADEF"))],
                  reverse=True)
    return os.path.join(repo_dir, cand[0]) if cand else None

# distribuzioni FX di ripiego (skill, se manca il file)
FALLBACK = {
    "DELTA": {"EUR": .5517, "USD": .2961, "GBP": .0924, "CHF": .0450, "NOK": .0148},
    "DELTADEF": {"EUR": .60, "USD": .27, "GBP": .07, "CHF": .04, "NOK": .02},
    "SUPERDISCOVERY": {"Other EM": .38, "KRW": .15, "CNY": .14, "EUR": .13, "TWD": .10, "INR": .04, "BRL": .04, "HKD": .02},
    "ALPHAGREEN": {"USD": .81, "EUR": .12, "HKD": .03, "CNY": .02, "DKK": .02},
}

GRAN_CCYS = ["EUR","USD","GBP","CHF","JPY","NOK","DKK","KRW","TWD","CNY","HKD","INR","BRL","MXN","IDR"]
SYN_CCYS  = ["EUR","USD","GBP","CHF","JPY","NOK","KRW","CNY","INR","BRL","Other EM"]
_MINORI   = ["DKK","TWD","HKD","MXN","IDR","SEK","CAD","THB","CLP","ZAR","PLN","HUF","COP","PEN","EGP","QAR","AED","SAR","TRY"]
AZ_LABEL  = {"USD":"Azioni dirette USA (incl. Digital Realty)","EUR":"Azioni dirette Europa Eurozone",
             "GBP":"Azioni dirette UK","CHF":"Azioni dirette Svizzera","KRW":"Azione Samsung Electronics GDR (KRW)",
             "DKK":"Azione Novo Nordisk (DKK)","JPY":"Azione Komatsu (JPY)"}

def _fund_dist(repo_dir, tipo, holding):
    fp = _find_fund_file(repo_dir, tipo)
    if fp:
        fr = fund_fx_fractions(fp, tipo); src = os.path.basename(fp)
    else:
        fr = FALLBACK.get(tipo, {"EUR": 1.0}); src = "ripiego (skill)"
    return {ccy: f * holding for ccy, f in fr.items()}, src

def build_matrix(pf, repo_dir):
    """Matrice valutaria look-through. Struttura come esempio Camperio:
    gran_rows (dettaglio granulare), syn_rows (3 macro + Other EM/Altro), oro, derivati."""
    nav = float(pf["meta"]["nav"])
    az = defaultdict(float); bond = defaultdict(float); cash = defaultdict(float)
    ieem = 0.0; mdax = 0.0; oro = 0.0; hold = {}
    for p in pf["positions"]:
        g = (p.get("grutit") or "").upper(); cod = p.get("codabi") or ""; v = p.get("valmer", 0.0)
        if cod == "RATEI": cash["EUR"] += v; continue
        if g == "H19": oro += v; continue
        if cod in FUND_CODE: hold[FUND_CODE[cod]] = v; continue
        if cod == "0M6317": ieem += v; continue
        if cod == "E27585": mdax += v; continue
        if g.startswith("A"):
            c = M.currency_of(p) or "EUR"; bond[M.CCY_LABEL.get(c, c)] += v; continue
        if g.startswith("Z"):
            c = M.currency_of(p) or "EUR"; cash[M.CCY_LABEL.get(c, c)] += v; continue
        c = M.currency_of(p)
        if c is None: oro += v; continue
        c = M.CCY_LABEL.get(c, c)
        if c == "Other EM": ieem += v
        else: az[c] += v
    fonti = {}
    delta, fonti["DELTA"] = _fund_dist(repo_dir, "DELTA", hold.get("DELTA", 0))
    sup, fonti["SUPERDISCOVERY"] = _fund_dist(repo_dir, "SUPERDISCOVERY", hold.get("SUPERDISCOVERY", 0))
    alpha, fonti["ALPHAGREEN"] = _fund_dist(repo_dir, "ALPHAGREEN", hold.get("ALPHAGREEN", 0))
    deltadef, fonti["DELTADEF"] = _fund_dist(repo_dir, "DELTADEF", hold.get("DELTADEF", 0))

    H = "__header__"
    gran = [(H, "AZIONI DIRETTE", None)]
    for c in [x for x in GRAN_CCYS if x in az] + sorted(k for k in az if k not in GRAN_CCYS):
        gran.append((None, AZ_LABEL.get(c, "Azioni dirette " + c), {c: az[c]}))
    gran.append((H, "ETF AZIONARI", None))
    gran.append((None, "ETF MSCI EM (IEEM, paniere EM)", {"Other EM": ieem}))
    gran.append((None, "ETF MDAX (DE)", {"EUR": mdax}))
    gran.append((H, "OBBLIGAZIONI E CASH", None))
    for c in [x for x in GRAN_CCYS if x in bond] + sorted(k for k in bond if k not in GRAN_CCYS):
        gran.append((None, "Bond " + c, {c: bond[c]}))
    gran.append((None, "Cash & margini", dict(cash)))
    gran.append((H, "FONDI UCITS (look-through FX)", None))
    gran.append((None, "Fondo DELTA UCITS - look-through", delta))
    gran.append((None, "Fondo Superdiscovery - look-through", sup))
    gran.append((None, "Fondo Alpha Green - look-through", alpha))
    if hold.get("DELTADEF", 0):
        gran.append((None, "Fondo Delta Defensive - look-through", deltadef))

    # aggregati per valuta (esclude Oro)
    tot = defaultdict(float)
    for d in (az, bond, cash, delta, sup, alpha, deltadef):
        for c, x in d.items(): tot[c] += x
    tot["EUR"] += mdax; tot["Other EM"] += ieem

    # sintetica: 3 macro
    macro_az = defaultdict(float)
    for d in (az, delta, sup, alpha, deltadef):
        for c, x in d.items(): macro_az[c] += x
    macro_az["EUR"] += mdax; macro_az["Other EM"] += ieem
    macro_bond = defaultdict(float)
    for d in (bond, cash):
        for c, x in d.items(): macro_bond[c] += x

    def _syn(d):
        out = {c: d.get(c, 0.0) for c in SYN_CCYS}
        out["Altro (<1%)"] = sum(d.get(c, 0.0) for c in _MINORI)
        return out
    syn = [("Azioni, ETF e fondi (look-through)", _syn(macro_az)),
           ("Obbligazioni e liquidita", _syn(macro_bond)),
           ("Oro fisico", {"Oro": oro})]
    syn_tot = _syn(tot); syn_tot["Oro"] = oro

    # derivati informativi (non FX)
    der = []
    spf = next((x for x in pf.get("pod", []) if x.get("codabi") == "E35126" or "S&P" in str(x.get("des", ""))), None)
    if spf:
        der.append(("Future S&P 500 short (hedge equity)", {"USD": -abs(spf.get("pnet", 0)) * (spf.get("val") or 0) * 50 / 1.1358}))
    callu = next((p for p in pf["positions"] if p.get("codabi") == "E35492"), None)
    if callu:
        der.append(("il call US Ultra 10Y (hedge tassi USD)", {"USD": callu.get("valorefut", 0)}))

    return {"nav": nav, "oro": oro, "fonti": fonti, "gran": gran, "gran_tot": dict(tot),
            "syn": syn, "syn_tot": syn_tot, "derivati": der,
            "azioni": dict(macro_az), "bond": dict(macro_bond), "totale": dict(tot)}


# ==================== REPORT TITOLI LORDO/NETTO (single-name look-through) ====================
import csv, re
_STOP = set("INC CORP CORPORATION PLC AG SA NV LTD LIMITED HOLDING HOLDINGS HLDGS CO COMPANY GROUP GRP THE SE SPA CLASS CL REG REGD ADR GDR SPON SPONS SPONSORED ORD SAB CV PS PAR PARTECIPATION PFD PREF PRF SHS NPV BR RG N AB ASA OYJ".split())
_ALIAS = {"ALPHABET":"ALPHABET", "GOOGLE":"ALPHABET"}
def _norm(s):
    s = re.sub(r"[^A-Za-z0-9 ]", " ", str(s).upper())
    toks = [t for t in s.split() if t not in _STOP and len(t) > 1 and t not in ("A","B","C","DE")]
    k = " ".join(toks[:2])
    return _ALIAS.get(k, k)

def parse_ishares(path):
    """Ritorna [(nome, peso_frazione)] delle partecipazioni azionarie di un CSV iShares."""
    txt = open(path, encoding="utf-8-sig").read().splitlines()
    try:
        st = next(i for i, l in enumerate(txt) if l.startswith("Ticker"))
    except StopIteration:
        return []
    out = []
    for row in csv.reader(txt[st + 1:]):
        if len(row) < 6 or row[3].strip() != "Azionario":
            continue
        try:
            w = float(row[5].replace(".", "").replace(",", ".").replace("%", "")) / 100
        except Exception:
            continue
        out.append((row[1].strip(), w))
    return out

def _find_csv(repo_dir, prefix):
    if not os.path.isdir(repo_dir): return None
    c = sorted([f for f in os.listdir(repo_dir) if f.upper().startswith(prefix) and f.lower().endswith(".csv")], reverse=True)
    return os.path.join(repo_dir, c[0]) if c else None

# classificazione temi AI (temi-ai.md)
_TEMI = [
    ("Semiconduttori", ["NVIDIA","MICRON","INFINEON","ADVANCED MICRO","INTEL","BROADCOM","TAIWAN SEMICONDUCTOR","ASML","APPLIED MATERIALS","SAMSUNG ELECTRONICS","SK HYNIX"]),
    ("Hyperscalers", ["APPLE","MICROSOFT","ALPHABET","META PLATFORMS","AMAZON"]),
    ("AI Software", ["PALANTIR","SALESFORCE","ORACLE","ADOBE","SERVICENOW","SNOWFLAKE"]),
    ("AI Infrastructure", ["SIEMENS","SCHNEIDER","CISCO","PRYSMIAN","ENEL","VERTIV","EATON","DIGITAL REALTY"]),
]
def _tema(name):
    u = name.upper()
    for t, kws in _TEMI:
        if any(k in u for k in kws): return t
    return ""

_PHYS = {"Bond Governative","Bond","Cash","Formation Expenses","Variation Margins","Committed Liquidity","Equity","ETF"}
def _fund_nav(rows):
    return sum((r[C_ESP] or 0) for r in rows if isinstance(r[C_ESP], numbers.Number) and r[C_TIPS] in _PHYS)

# DELTA: indice -> ETF proxy
_IDX_ETF = {"SPX": "IUSA", "SX5E": "EUE", "DAX": "EXS1", "UKX": "ISF", "SMI": "EXI1"}

def build_titoli(pf, repo_dir):
    """Costruisce la vista titoli single-name lordo/netto vs benchmark.
    Ritorna {nav, rows[], temi[], note}."""
    nav = float(pf["meta"]["nav"])
    agg = {}
    def A(name, bucket, val, src=None):
        k = _norm(name)
        if not k: return
        d = agg.setdefault(k, {"name": name.title(), "diretto": 0.0, "fondi": 0.0, "indici": 0.0,
                               "hedge": 0.0, "src": set()})
        d[bucket] += val
        if src: d["src"].add(src)

    # holdings ETF (pesi)
    def hold(tk):
        p = _find_csv(repo_dir, tk)
        return parse_ishares(p) if p else []
    H = {tk: hold(tk) for tk in ("IUSA", "EUE", "EXS1", "ISF", "EXI1", "IEEM", "EUN", "EXS3")}

    # 1) DIRETTO + identifica ETF e fondi detenuti
    FUND = {"95ZYC2": "DELTA", "8W8C01": "SUPERDISCOVERY", "54OPLF": "ALPHAGREEN",
            "8W8C34": "DELTADEF", "8W8C45": "DELTADEF", "50QNF8": "DELTADEF", "52PZG8": "DELTADEF"}
    ETFD = {"0M6317": "IEEM", "E27585": "EXS3"}
    fund_hold = {}; etf_hold = {}
    for p in pf["positions"]:
        g = (p.get("grutit") or "").upper(); cod = p["codabi"]; v = p["valmer"]; des = p["des"]
        if cod in FUND: fund_hold[FUND[cod]] = v; continue
        if cod in ETFD: etf_hold[ETFD[cod]] = v; continue
        if g.startswith("E") or g.startswith("B") or g == "H23":
            A(des, "diretto", v, "Diretto")

    # 2) FONDI: equity diretta + DELTA indici + MXEF
    def frows(prefix):
        fp = _find_fund_file(repo_dir, prefix)
        if not fp: return None
        try:
            import openpyxl
            wb = openpyxl.load_workbook(fp, data_only=True, read_only=True)
            return list(wb[wb.sheetnames[0]].iter_rows(min_row=6, values_only=True))
        except Exception:
            return None
    # Superdiscovery
    sd = frows("SUPERDISCOVERY")
    if sd and fund_hold.get("SUPERDISCOVERY"):
        share = fund_hold["SUPERDISCOVERY"] / (_fund_nav(sd) or 1)
        mxef = 0.0
        for r in sd:
            t = r[C_TIPS]; e = r[C_ESP] if isinstance(r[C_ESP], numbers.Number) else 0
            if t in ("Equity", "ETF"): A(r[C_DESC], "fondi", e * share, "Superdiscovery")
            elif t == "Future": mxef += e
        for nome, w in H["IEEM"]: A(nome, "fondi", mxef * share * w, "MXEF(Disc)")
    # Alpha
    al = frows("ALPHAGREEN")
    if al and fund_hold.get("ALPHAGREEN"):
        share = fund_hold["ALPHAGREEN"] / (_fund_nav(al) or 1)
        for r in al:
            if r[C_TIPS] == "Equity": A(r[C_DESC], "fondi", (r[C_ESP] or 0) * share, "Alpha Green")
    # DELTA indici
    dl = frows("DELTA")
    if dl and fund_hold.get("DELTA"):
        share = fund_hold["DELTA"] / (_fund_nav(dl) or 1)
        idx = defaultdict(float)
        for r in dl:
            if r[C_TIPS] in ("Option", "Future") and (r[C_MSEC] or "") == "Index":
                key = str(r[65] or r[64] or "")
                for k, nm in [("SX5E","SX5E"),("DAX","DAX"),("UKX","UKX"),("SMI","SMI"),("ES","SPX"),("SPX","SPX")]:
                    if k in key: idx[nm] += (r[C_ESP] or 0); break
        for ix, esp in idx.items():
            etf = _IDX_ETF.get(ix)
            for nome, w in H.get(etf, []):
                A(nome, "indici", esp * share * w, ix + "(DELTA)")
    # DELTA DEFENSIVE indici (stessa logica del DELTA)
    ddf = frows("DELTADEF")
    if ddf and fund_hold.get("DELTADEF"):
        share = fund_hold["DELTADEF"] / (_fund_nav(ddf) or 1)
        idx = defaultdict(float)
        for r in ddf:
            if r[C_TIPS] in ("Option", "Future") and (r[C_MSEC] or "") == "Index":
                key = str(r[65] or r[64] or "")
                for k, nm in [("SX5E","SX5E"),("DAX","DAX"),("UKX","UKX"),("SMI","SMI"),("ES","SPX"),("SPX","SPX")]:
                    if k in key: idx[nm] += (r[C_ESP] or 0); break
        for ix, esp in idx.items():
            etf = _IDX_ETF.get(ix)
            for nome, w in H.get(etf, []):
                A(nome, "indici", esp * share * w, ix + "(DeltaDef)")
    # 3) ETF detenuti direttamente (look-through)
    for nome, w in H["IEEM"]: A(nome, "indici", etf_hold.get("IEEM", 0) * w, "ETF IEEM")
    for nome, w in H["EXS3"]: A(nome, "indici", etf_hold.get("EXS3", 0) * w, "ETF MDAX")

    # 4) HEDGE S&P short del portafoglio
    spf = next((x for x in pf.get("pod", []) if "S&P" in str(x.get("des", "")) or x.get("codabi") == "E35126"), None)
    notional = abs(spf["pnet"]) * spf["val"] * 50 / 1.1358 if spf else 0.0
    for nome, w in H["IUSA"]: A(nome, "hedge", -notional * w, None)

    # 5) BENCHMARK titoli (20% IUSA + 20% EUN + 10% IEEM)
    bench = defaultdict(float)
    for nome, w in H["IUSA"]: bench[_norm(nome)] += 0.20 * w
    for nome, w in H["EUN"]: bench[_norm(nome)] += 0.20 * w
    for nome, w in H["IEEM"]: bench[_norm(nome)] += 0.10 * w

    rows = []
    for k, d in agg.items():
        lordo = d["diretto"] + d["fondi"] + d["indici"]
        if abs(lordo) < 1 and abs(d["hedge"]) < 1: continue
        netto = lordo + d["hedge"]
        bw = bench.get(k, 0.0)
        rows.append({"name": d["name"], "fonte": " + ".join(sorted(d["src"])) or "Diretto",
                     "tema": _tema(d["name"]), "lordo": lordo, "pct_lordo": lordo / nav,
                     "hedge": d["hedge"], "netto": netto, "pct_netto": netto / nav,
                     "bench": bw, "delta": netto / nav - bw})
    rows.sort(key=lambda x: -x["lordo"])
    # temi AI
    temi = {}
    for t, _ in _TEMI:
        lg = sum(r["lordo"] for r in rows if r["tema"] == t)
        nt = sum(r["netto"] for r in rows if r["tema"] == t)
        bn = sum(r["bench"] for r in rows if r["tema"] == t)
        temi[t] = {"lordo": lg / nav, "netto": nt / nav, "bench": bn}
    tot_dir = sum(d["diretto"] for d in agg.values())
    tot_fon = sum(d["fondi"] for d in agg.values())
    tot_idx = sum(d["indici"] for d in agg.values())
    tot_hed = sum(d["hedge"] for d in agg.values())
    return {"nav": nav, "rows": rows, "temi": temi,
            "n": len(rows), "hedge_notional": notional, "fonti": list(fund_hold.keys()),
            "tot_diretto": tot_dir, "tot_fondi": tot_fon, "tot_indici": tot_idx,
            "tot_hedge": tot_hed, "tot_netto": tot_dir + tot_fon + tot_idx + tot_hed}

# ==================== RENDERER HTML / EXCEL ====================
from camperio_core.render.format import eur, pct
_CCY_ORDER = ["EUR","USD","GBP","CHF","JPY","NOK","DKK","SEK","KRW","TWD","CNY","HKD","INR","BRL","MXN","Other EM"]
def _ccy_sorted(keys):
    extra = [k for k in keys if k not in _CCY_ORDER]
    return [k for k in _CCY_ORDER if k in keys] + sorted(extra)

def matrice_html(m):
    nav = m["nav"]
    cols = SYN_CCYS + ["Altro (<1%)", "Oro"]
    head = ('<div class="rh"><div class="rt">Matrice Valutaria - look-through</div>'
            '<div class="rs">% NAV - vista sintetica (dettaglio completo nell\'Excel) - Samsung GDR&#8594;KRW, ETF EM&#8594;Other EM, oro fuori FX</div></div>')
    th = "".join(f"<th class=r>{c}</th>" for c in cols)
    def rrow(lbl, d, cls=""):
        cells = "".join(f"<td class=r>{pct(d.get(c,0)/nav*100,2,seg=False) if abs(d.get(c,0))>1 else ''}</td>" for c in cols)
        return f"<tr class='{cls}'><td><b>{lbl}</b></td>{cells}</tr>"
    body = "".join(rrow(l, d) for l, d in m["syn"]) + rrow("TOTALE", m["syn_tot"], "tot")
    fonti = " - ".join(f"{k}: {v}" for k, v in m["fonti"].items())
    totfx = sum(v for c, v in m["gran_tot"].items())
    note = (f'<div class="note">Somma FX (escl. oro): {pct(totfx/nav*100,1,seg=False)} NAV - Oro {pct(m["oro"]/nav*100,2,seg=False)} (fuori FX). '
            f'Look-through fondi: {fonti}.</div>')
    return (head + '<h3>Matrice sintetica per tipologia (% NAV)</h3><div style="overflow-x:auto">'
            '<table><thead><tr><th>Tipologia</th>' + th + '</tr></thead><tbody>'
            + body + '</tbody></table></div>' + note)

def _xl_styles():
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    BLU = "151F6D"
    return {"hfill": PatternFill("solid", fgColor=BLU), "hfont": Font(color="FFFFFF", bold=True, size=9.5),
            "sec": PatternFill("solid", fgColor="2E5A9E"), "secf": Font(color="FFFFFF", bold=True, size=9),
            "totf": PatternFill("solid", fgColor="E9E6DC"), "alt": PatternFill("solid", fgColor="F7F5EF"),
            "title": Font(bold=True, size=13, color=BLU), "sub": Font(size=9, color="666666"),
            "ital": Font(italic=True, color="C9A66B"),
            "border": Border(bottom=Side(style="thin", color="E0DDD4"))}

def matrice_excel(m, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    from openpyxl.utils import get_column_letter
    nav = m["nav"]; st = _xl_styles()
    wb = Workbook()

    # ---- COMPLETA (granulare) ----
    ws = wb.active; ws.title = "Matrice completa (%)"
    ccys = GRAN_CCYS
    ws["A1"] = "Matrice Strumento x Valuta - Linea Camperio"; ws["A1"].font = st["title"]
    ws["A2"] = "Look-through fondi - % NAV - Samsung GDR->KRW - ETF EM->Other EM (in foglio sintetica) - oro fuori FX"; ws["A2"].font = st["sub"]
    hdr = ["Strumento"] + ccys
    for j, h in enumerate(hdr, 1):
        c = ws.cell(row=4, column=j, value=h); c.fill = st["hfill"]; c.font = st["hfont"]; c.alignment = Alignment(horizontal="center" if j > 1 else "left")
    r = 5
    for kind, lbl, d in m["gran"]:
        if kind == "__header__":
            cc = ws.cell(row=r, column=1, value=lbl); cc.fill = st["sec"]; cc.font = st["secf"]
            for j in range(2, len(ccys) + 2): ws.cell(row=r, column=j).fill = st["sec"]
            r += 1; continue
        ws.cell(row=r, column=1, value=lbl).border = st["border"]
        for j, cc in enumerate(ccys, 2):
            v = d.get(cc, 0) / nav
            cell = ws.cell(row=r, column=j, value=round(v, 4) if abs(d.get(cc, 0)) > 1 else None)
            cell.number_format = "0.00%"; cell.border = st["border"]
        r += 1
    # TOTALE
    ws.cell(row=r, column=1, value="TOTALE ESPOSIZIONE PER VALUTA").font = Font(bold=True)
    ws.cell(row=r, column=1).fill = st["totf"]
    for j, cc in enumerate(ccys, 2):
        cell = ws.cell(row=r, column=j, value=round(m["gran_tot"].get(cc, 0) / nav, 4) or None)
        cell.number_format = "0.00%"; cell.font = Font(bold=True); cell.fill = st["totf"]
    rfx = r + 1
    ws.cell(row=rfx, column=1, value="% FX (escluso Oro)").font = Font(bold=True, color="666666")
    base = nav - m["oro"]
    for j, cc in enumerate(ccys, 2):
        cell = ws.cell(row=rfx, column=j, value=round(m["gran_tot"].get(cc, 0) / base, 4) or None)
        cell.number_format = "0.00%"; cell.font = Font(color="666666")
    r = rfx + 2
    ws.cell(row=r, column=1, value="Oro fisico (asset class autonoma, fuori FX): " + str(round(m["oro"]/nav*100, 2)) + "% NAV").font = st["ital"]; r += 1
    ws.cell(row=r, column=1, value="DERIVATI ON-BOOK (informativi - non impattano la matrice FX)").font = Font(bold=True, size=9); r += 1
    for lbl, d in m["derivati"]:
        ws.cell(row=r, column=1, value=lbl)
        for j, cc in enumerate(ccys, 2):
            if cc in d:
                cell = ws.cell(row=r, column=j, value=round(d[cc] / nav, 4)); cell.number_format = "0.00%"
        r += 1
    ws.column_dimensions["A"].width = 40; ws.freeze_panes = "B5"
    for j in range(2, len(ccys) + 2): ws.column_dimensions[get_column_letter(j)].width = 8

    # ---- SINTETICA ----
    ws2 = wb.create_sheet("Matrice sintetica (%)")
    scols = SYN_CCYS + ["Altro (<1%)", "Oro", "TOT %"]
    ws2["A1"] = "Matrice Valutaria sintetica - Linea Camperio"; ws2["A1"].font = st["title"]
    ws2["A2"] = "% NAV - vista aggregata look-through (fondi+ETF+azioni)"; ws2["A2"].font = st["sub"]
    for j, h in enumerate(["Tipologia"] + scols, 1):
        c = ws2.cell(row=4, column=j, value=h); c.fill = st["hfill"]; c.font = st["hfont"]; c.alignment = Alignment(horizontal="center" if j > 1 else "left")
    r = 5
    def _rowsum(d): return sum(d.get(c, 0) for c in SYN_CCYS + ["Altro (<1%)", "Oro"])
    for lbl, d in m["syn"]:
        ws2.cell(row=r, column=1, value=lbl).border = st["border"]
        for j, cc in enumerate(scols, 2):
            val = (_rowsum(d) if cc == "TOT %" else d.get(cc, 0)) / nav
            cell = ws2.cell(row=r, column=j, value=round(val, 4) if abs(val * nav) > 1 else None)
            cell.number_format = "0.00%"; cell.border = st["border"]
        r += 1
    ws2.cell(row=r, column=1, value="TOTALE").font = Font(bold=True); ws2.cell(row=r, column=1).fill = st["totf"]
    for j, cc in enumerate(scols, 2):
        val = (_rowsum(m["syn_tot"]) if cc == "TOT %" else m["syn_tot"].get(cc, 0)) / nav
        cell = ws2.cell(row=r, column=j, value=round(val, 4) or None); cell.number_format = "0.00%"; cell.font = Font(bold=True); cell.fill = st["totf"]
    r += 2
    for txt in ["'Altro (<1%)' = DKK + TWD + HKD + MXN + IDR e altre minori.",
                "Vista FX fondi: DELTA % nette factsheet; Superdiscovery per paese (USD residuo->Other EM); Alpha valuta reale.",
                "Oro: asset class autonoma, fuori dal computo FX."]:
        ws2.cell(row=r, column=1, value=txt).font = st["sub"]; r += 1
    ws2.column_dimensions["A"].width = 34
    for j in range(2, len(scols) + 2): ws2.column_dimensions[get_column_letter(j)].width = 9
    wb.save(path); return path

def titoli_html(t):
    nav = t["nav"]
    head = ('<div class="rh"><div class="rt">Titoli azionari — look-through lordo/netto vs benchmark</div>'
            f'<div class="rs">{t["n"]} nomi · lordo = diretto + fondi + indici (DELTA delta-adj, MXEF→IEEM) · netto = lordo − hedge S&P · bench 20% IUSA+20% EUN+10% IEEM</div></div>')
    # temi AI
    kp = ""
    for tema, v in t["temi"].items():
        col = "1F7A4D" if (v["lordo"] - v["bench"]) >= 0 else "B3261E"
        kp += (f'<div class="ki"><div class="kl">{tema}</div>'
               f'<div class="kv" style="color:#{col};font-size:18px">{pct(v["lordo"]*100)}</div>'
               f'<div class="ks">bench {pct(v["bench"]*100,2,seg=False)} · Δ {pct((v["lordo"]-v["bench"])*100)} p.p.</div></div>')
    rows = ""
    for i, r in enumerate(t["rows"][:30], 1):
        dcol = "1F7A4D" if r["delta"] >= 0 else "B3261E"
        rows += (f"<tr><td class=r>{i}</td><td><b>{r['name']}</b></td><td style='font-size:11px'>{r['fonte']}</td>"
                 f"<td>{r['tema']}</td><td class=r>{eur(r['lordo'])}</td><td class=r>{pct(r['pct_lordo']*100,2,seg=False)}</td>"
                 f"<td class=r>{eur(r['hedge']) if r['hedge'] else '—'}</td><td class=r>{eur(r['netto'])}</td>"
                 f"<td class=r>{pct(r['pct_netto']*100,2,seg=False)}</td><td class=r>{pct(r['bench']*100,2,seg=False)}</td>"
                 f"<td class=r style='color:#{dcol}'>{pct(r['delta']*100)}</td></tr>")
    return (head + '<h3>Temi AI (lordo vs benchmark)</h3><div class="kp">' + kp + '</div>'
            + f'<h3>Primi 30 titoli (su {t["n"]}) — la tabella completa è nell\'Excel</h3>'
            + '<div style="overflow-x:auto"><table><thead><tr><th class=r>#</th><th>Titolo</th><th>Fonte</th><th>Tema AI</th>'
            + '<th class=r>€ lordo</th><th class=r>% lordo</th><th class=r>Hedge S&P</th><th class=r>€ netto</th>'
            + '<th class=r>% netto</th><th class=r>% bench</th><th class=r>Δ p.p.</th></tr></thead><tbody>'
            + rows + '</tbody></table></div>')

def titoli_excel(t, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    BLU = "151F6D"
    hfill = PatternFill("solid", fgColor=BLU); hfont = Font(color="FFFFFF", bold=True, size=9)
    alt = PatternFill("solid", fgColor="F4F2EC")
    thin = Side(style="thin", color="E0DDD4"); border = Border(bottom=thin)
    wb = Workbook(); ws = wb.active; ws.title = "Look-through lordo-netto"
    ws["A1"] = "Titoli azionari - look-through LORDO/NETTO vs benchmark - Linea Camperio"
    ws["A1"].font = Font(bold=True, size=13, color=BLU)
    ws["A2"] = ("Lordo = diretto + equity fondi + indici fondi (DELTA delta-adj: S&P->IUSA, EuroStoxx->EUE, DAX->EXS1, "
                "FTSE->ISF, SMI->EXI1; Superdiscovery MXEF->IEEM). Netto = lordo - hedge S&P short. Bench = 20% IUSA + 20% EUN + 10% IEEM.")
    ws["A2"].font = Font(size=8, color="666666")
    hdr = ["#", "Titolo", "Fonte", "Tema AI", "EUR lordo", "% NAV lordo", "Hedge S&P EUR", "EUR netto", "% NAV netto", "% bench", "Delta vs bench p.p."]
    for j, h in enumerate(hdr, 1):
        c = ws.cell(row=4, column=j, value=h); c.fill = hfill; c.font = hfont; c.alignment = Alignment(horizontal="center", wrap_text=True)
    r = 5
    for i, x in enumerate(t["rows"], 1):
        vals = [i, x["name"], x["fonte"], x["tema"], round(x["lordo"], 0), x["pct_lordo"],
                round(x["hedge"], 0) or None, round(x["netto"], 0), x["pct_netto"], x["bench"], x["delta"]]
        for j, v in enumerate(vals, 1):
            cell = ws.cell(row=r, column=j, value=v); cell.border = border
            if j in (5, 7, 8): cell.number_format = "#,##0"
            if j in (6, 9, 10): cell.number_format = "0.00%"
            if j == 11: cell.number_format = "+0.00%;-0.00%"
            if r % 2 == 0: cell.fill = alt
        r += 1
    ws.freeze_panes = "A5"; ws.auto_filter.ref = "A4:K" + str(r - 1)
    for col, w in zip("ABCDEFGHIJK", (5, 30, 30, 16, 12, 11, 13, 12, 11, 9, 13)):
        ws.column_dimensions[col].width = w
    # foglio Temi AI
    ws2 = wb.create_sheet("Temi AI")
    ws2["A1"] = "Esposizioni tematiche AI - lordo / netto vs benchmark"; ws2["A1"].font = Font(bold=True, size=13, color=BLU)
    for j, h in enumerate(["Tema", "% NAV lordo", "% NAV netto", "% benchmark", "Delta lordo p.p."], 1):
        c = ws2.cell(row=3, column=j, value=h); c.fill = hfill; c.font = hfont
    rr = 4
    for tema, v in t["temi"].items():
        ws2.cell(row=rr, column=1, value=tema)
        for j, val in enumerate([v["lordo"], v["netto"], v["bench"], v["lordo"] - v["bench"]], 2):
            cell = ws2.cell(row=rr, column=j, value=val); cell.number_format = "0.00%"
        rr += 1
    ws2.column_dimensions["A"].width = 20
    for col in "BCDE": ws2.column_dimensions[col].width = 13
    wb.save(path); return path


# ---- Word compatti (branding minimo) ----
_LOGO = os.path.join(HERE, "static", "logo-camperio.png")
_LOGO_W = os.path.join(HERE, "static", "logo-bianco.png")
_F1 = "Via Camperio, 9 — 20123 Milano · Tel +39 02.50020918 · camperioSIM@camperiosim.com · www.camperiosim.com"
_F2 = "Consob — delibera n. 11761 del 22/12/1998 — albo n. 48. C.F. 02342760275 — P.IVA 11791000158 — Cod. Banca d'Italia 16206/5 — FNG SIM0077."

def _doc(titolo, sub):
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    d = Document(); s = d.sections[0]
    s.top_margin = Mm(30); s.bottom_margin = Mm(18); s.left_margin = Mm(16); s.right_margin = Mm(16)
    hdr = s.header
    if hdr.paragraphs and hdr.paragraphs[0].runs == []:
        hdr.paragraphs[0].text = ""
    tb = hdr.add_table(rows=1, cols=2, width=Mm(178))
    try: tb.columns[0].width = Mm(120); tb.columns[1].width = Mm(58)
    except Exception: pass
    for cc in tb.rows[0].cells:
        tcPr = cc._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "151F6D"); tcPr.append(sh)
    cL, cR = tb.rows[0].cells
    try: cL.paragraphs[0].add_run().add_picture(_LOGO_W, width=Mm(46))
    except Exception: pass
    pR = cR.paragraphs[0]; pR.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pR.add_run("USO INTERNO · COMITATO"); rr.bold = True; rr.font.size = Pt(8.5); rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p = d.add_paragraph(); r = p.add_run(titolo); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x15,0x1F,0x6D)
    q = d.add_paragraph(); rq = q.add_run(sub); rq.font.size = Pt(9); rq.font.color.rgb = RGBColor(0x66,0x66,0x66)
    fp = s.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(_F1 + "\n" + _F2); fr.font.size = Pt(6.3); fr.font.color.rgb = RGBColor(0x66,0x66,0x66)
    return d

def _table(d, headers, rows):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    t = d.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cc = t.rows[0].cells[i]; cc.text = ""; rn = cc.paragraphs[0].add_run(h); rn.bold = True
        rn.font.size = Pt(8); rn.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tcPr = cc._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), "151F6D"); tcPr.append(sh)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; rn = cells[i].paragraphs[0].add_run(str(v)); rn.font.size = Pt(8)

def matrice_word(m, path):
    nav = m["nav"]; tot = m["totale"]; ccys = _ccy_sorted(set(list(m["azioni"]) + list(m["bond"])))
    d = _doc("Matrice Valutaria — Linea Camperio", f"% NAV · look-through fondi · NAV {eur(nav)}")
    def rr(lbl, dd): return [lbl] + [pct(dd.get(c,0)/nav*100,2,seg=False) if dd.get(c,0) else "—" for c in ccys]
    _table(d, ["Componente"] + ccys, [rr("Azioni/ETF/fondi", m["azioni"]), rr("Obblig. e liquidità", m["bond"]), rr("TOTALE FX", tot)])
    d.add_paragraph(f"Oro fisico (fuori FX): {eur(m['oro'])} = {pct(m['oro']/nav*100,2,seg=False)} NAV.")
    d.save(path); return path

def titoli_word(t, path):
    d = _doc("Titoli azionari — look-through lordo/netto vs benchmark", f"{t['n']} nomi · diretto + fondi + indici (delta-adj) − hedge S&P · bench 20/20/10")
    d.add_paragraph().add_run("Temi AI").bold = True
    _table(d, ["Tema", "% lordo", "% netto", "% bench", "Δ lordo p.p."],
           [[k, pct(v["lordo"]*100), pct(v["netto"]*100), pct(v["bench"]*100,2,seg=False), pct((v["lordo"]-v["bench"])*100)] for k, v in t["temi"].items()])
    d.add_paragraph().add_run(f"Primi 40 titoli (completo nell'Excel)").bold = True
    _table(d, ["#", "Titolo", "€ lordo", "% lordo", "€ netto", "% netto", "% bench", "Δ p.p."],
           [[i, x["name"], eur(x["lordo"]), pct(x["pct_lordo"]*100,2,seg=False), eur(x["netto"]),
             pct(x["pct_netto"]*100,2,seg=False), pct(x["bench"]*100,2,seg=False), pct(x["delta"]*100)]
            for i, x in enumerate(t["rows"][:40], 1)])
    d.save(path); return path


# ==================== REPORT VARIAZIONI PREZZI (azioni dirette, prezzo EUR/azione) ====================
# Settore GICS (inglese) dai CSV iShares (colonna "Settore" in italiano -> mappa)
_GICS_IT = {"IT": "Information Technology", "Industriali": "Industrials", "Finanziari": "Financials",
            "Salute": "Health Care", "Consumi Discrezionali": "Consumer Discretionary",
            "Generi di largo consumo": "Consumer Staples", "Imprese di servizi di pubblica utilità": "Utilities",
            "Immobili": "Real Estate", "Materiali": "Materials", "Comunicazione": "Communication Services",
            "Energia": "Energy"}
_SECT_FB = {"AENA": "Industrials", "ALSTOM": "Industrials", "EASYJET": "Consumer Discretionary",
            "ESSILOR": "Consumer Discretionary", "LABORATORY": "Health Care", "PHILIPS": "Health Care",
            "PRYSMIAN": "Industrials", "SCHLUMBERGER": "Energy", "THALES": "Industrials",
            "KOMATSU": "Industrials", "TECHNOGYM": "Consumer Discretionary"}
_SSTOP = set(("INC CORP CORPORATION CO LTD LIMITED PLC SA SE AG NV SPA THE GROUP HOLDING HOLDINGS COMPANY "
              "REG ORD PREF PFD SHS NPV CL CLASS ADR GDR SME HENNESSY LOUI SAS NEW OLD PAR HLD").split())
_SEXCH = set("FP GY IM LN NA US SW DC CHF USD EUR GBP DKK JPY SMD".split())

def _skeys(name):
    s = re.sub(r"\(.*?\)", " ", str(name or "").upper())
    s = re.sub(r"[^A-Z0-9 ]", " ", s)
    t = [x for x in s.split() if len(x) > 1 and x not in _SSTOP and x not in _SEXCH]
    ks = []
    if t: ks.append(t[0])
    if len(t) >= 2: ks.append(t[0] + " " + t[1])
    return ks

def _build_sector_map(repo_dir):
    smap = {}
    if not os.path.isdir(repo_dir): return smap
    for f in os.listdir(repo_dir):
        if not f.lower().endswith(".csv"): continue
        try:
            raw = open(os.path.join(repo_dir, f), encoding="utf-8-sig", errors="replace").read().splitlines()
            hi = next((i for i, l in enumerate(raw[:15]) if l.startswith("Ticker") and "Nome" in l), None)
            if hi is None: continue
            rd = list(csv.reader(raw[hi:])); h = rd[0]
            ni = h.index("Nome"); si = h.index("Settore")
            for r in rd[1:]:
                if len(r) > max(ni, si) and r[ni]:
                    g = _GICS_IT.get(r[si].strip(), r[si].strip())
                    for k in _skeys(r[ni]): smap.setdefault(k, g)
        except Exception:
            continue
    return smap

def _sector_of(name, smap):
    ks = _skeys(name)
    for k in reversed(ks):
        if k in smap: return smap[k]
    return _SECT_FB.get(ks[0], "n.d.") if ks else "n.d."

def _it_date(iso):
    try:
        y, m, d = iso.split("-"); return d + "/" + m + "/" + y
    except Exception:
        return iso

def build_variazioni(pf, repo_dir):
    """Variazioni settimanali di prezzo (EUR/azione) dei titoli azionari diretti, dal->al."""
    import data_layer as DL
    meta = pf["meta"]
    data = DL.price_changes(meta["schema"], meta["codcli"], meta.get("data_prec"), meta["data"])
    nav = float(data.get("nav") or meta.get("nav") or 0) or 1.0
    smap = _build_sector_map(repo_dir)
    rows = []
    for x in data.get("rows", []):
        ld, la, fd, fa = x.get("loc_dal"), x.get("loc_al"), x.get("fx_dal"), x.get("fx_al")
        e_dal = (ld / fd) if (ld and fd) else None
        e_al = (la / fa) if (la and fa) else None
        var = (e_al / e_dal - 1) if (e_dal and e_al and e_dal != 0) else None
        qa = x.get("qty_al", x.get("qty", 0))
        peso = (qa * e_al / nav) if (e_al and nav) else 0.0
        rows.append({"name": (x.get("nome") or "").title().strip(), "ccy": x.get("ccy", "EUR"),
                     "settore": _sector_of(x.get("nome"), smap), "eur_dal": e_dal, "eur_al": e_al,
                     "var": var, "peso": peso})
    valid = [r for r in rows if r["var"] is not None]
    valid.sort(key=lambda r: -r["var"])
    novar = [r for r in rows if r["var"] is None]
    up = sum(1 for r in valid if r["var"] > 0); down = sum(1 for r in valid if r["var"] < 0)
    media = (sum(r["var"] for r in valid) / len(valid)) if valid else 0.0
    wsum = sum(r["peso"] for r in valid) or 1.0
    media_w = (sum(r["var"] * r["peso"] for r in valid) / wsum) if valid else 0.0
    dal_iso = data.get("dal") or meta.get("data_prec"); al_iso = data.get("al") or meta.get("data")
    return {"nav": nav, "dal": _it_date(dal_iso), "al": _it_date(al_iso),
            "rows": valid + novar, "n": len(valid), "up": up, "down": down,
            "media": media, "media_w": media_w,
            "best": valid[0] if valid else None, "worst": valid[-1] if valid else None}

def _pe(x):
    if x is None: return "—"
    s = "{:,.2f}".format(x)
    return s.replace(",", "§").replace(".", ",").replace("§", ".")

def variazioni_html(d):
    head = ('<div class="rh"><div class="rt">Variazioni prezzi — azioni in portafoglio</div>'
            f'<div class="rs">Prezzo per azione in € · {d["dal"]} &#8594; {d["al"]} · var. include effetto cambio per i titoli in valuta · {d["n"]} titoli</div></div>')
    def card(lbl, val, sub=""):
        return f'<div class="ki"><div class="kl">{lbl}</div><div class="kv">{val}</div><div class="ks">{sub}</div></div>'
    mc = "1F7A4D" if d["media"] >= 0 else "B3261E"
    best = d["best"]; worst = d["worst"]
    kp = ('<div class="kp">'
          + card("Media semplice", f'<span style="color:#{mc}">{pct(d["media"]*100)}</span>')
          + card("Media ponderata", pct(d["media_w"]*100))
          + card("In rialzo / ribasso", f'{d["up"]} / {d["down"]}')
          + card("Migliore", (pct(best["var"]*100) if best else "—"), best["name"] if best else "")
          + card("Peggiore", (pct(worst["var"]*100) if worst else "—"), worst["name"] if worst else "")
          + '</div>')
    rr = ""
    for i, r in enumerate(d["rows"], 1):
        col = "1F7A4D" if (r["var"] or 0) >= 0 else "B3261E"
        vtxt = pct(r["var"]*100) if r["var"] is not None else "n.d."
        rr += (f"<tr><td class=r>{i}</td><td><b>{r['name']}</b></td><td>{r['settore']}</td>"
               f"<td>{r['ccy']}</td><td class=r>{_pe(r['eur_dal'])}</td><td class=r>{_pe(r['eur_al'])}</td>"
               f"<td class=r style='color:#{col}'>{vtxt}</td><td class=r>{pct(r['peso']*100,2,seg=False)}</td></tr>")
    return (head + '<h3>Sintesi settimana</h3>' + kp
            + '<h3>Dettaglio per titolo (ordinato per variazione)</h3>'
            + '<div style="overflow-x:auto"><table><thead><tr><th class=r>#</th><th>Titolo</th><th>Settore</th>'
            + '<th>Divisa</th><th class=r>Prezzo € ' + d["dal"][:5] + '</th><th class=r>Prezzo € ' + d["al"][:5] + '</th>'
            + '<th class=r>Var. %</th><th class=r>% NAV</th></tr></thead><tbody>' + rr + '</tbody></table></div>')

def variazioni_excel(d, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    BLU = "151F6D"
    hfill = PatternFill("solid", fgColor=BLU); hfont = Font(color="FFFFFF", bold=True, size=9)
    alt = PatternFill("solid", fgColor="F4F2EC"); border = Border(bottom=Side(style="thin", color="E0DDD4"))
    sub = Font(size=8, color="666666"); bold = Font(bold=True)
    wb = Workbook(); ws = wb.active; ws.title = "Variazioni prezzi"
    ws["A1"] = "Variazione settimanale dei prezzi — azioni in portafoglio"
    ws["A1"].font = Font(bold=True, size=13, color=BLU)
    ws["A2"] = ("Prezzo per azione in € · " + d["dal"] + " → " + d["al"]
                + " · var. include effetto cambio per i titoli in valuta · USO INTERNO"); ws["A2"].font = sub
    # sintesi
    r = 4
    ws.cell(row=r, column=1, value="Sintesi settimana").font = bold; r += 1
    syn = [("Media semplice", d["media"], "0.00%"), ("Media ponderata per peso", d["media_w"], "0.00%"),
           ("Titoli in rialzo", d["up"], "0"), ("Titoli in ribasso", d["down"], "0")]
    if d["best"]: syn.append(("Migliore", d["best"]["var"], "0.00%"))
    if d["worst"]: syn.append(("Peggiore", d["worst"]["var"], "0.00%"))
    for lbl, val, fmt in syn:
        ws.cell(row=r, column=1, value=lbl)
        c = ws.cell(row=r, column=2, value=val); c.number_format = fmt
        if lbl == "Migliore" and d["best"]: ws.cell(row=r, column=3, value=d["best"]["name"]).font = sub
        if lbl == "Peggiore" and d["worst"]: ws.cell(row=r, column=3, value=d["worst"]["name"]).font = sub
        r += 1
    r += 1
    hdr = ["#", "Titolo", "Settore", "Divisa", "Prezzo € " + d["dal"], "Prezzo € " + d["al"], "Var. % settimana", "% NAV"]
    hrow = r
    for j, h in enumerate(hdr, 1):
        c = ws.cell(row=hrow, column=j, value=h); c.fill = hfill; c.font = hfont
        c.alignment = Alignment(horizontal="center" if j > 1 else "left", wrap_text=True)
    r += 1
    for i, x in enumerate(d["rows"], 1):
        vals = [i, x["name"], x["settore"], x["ccy"],
                round(x["eur_dal"], 4) if x["eur_dal"] is not None else None,
                round(x["eur_al"], 4) if x["eur_al"] is not None else None,
                x["var"], x["peso"]]
        for j, v in enumerate(vals, 1):
            c = ws.cell(row=r, column=j, value=v); c.border = border
            if j in (5, 6): c.number_format = "#,##0.00"
            if j == 7: c.number_format = "+0.00%;-0.00%"
            if j == 8: c.number_format = "0.00%"
            if r % 2 == 0: c.fill = alt
        r += 1
    ws.freeze_panes = "A" + str(hrow + 1); ws.auto_filter.ref = "A" + str(hrow) + ":H" + str(r - 1)
    for col, w in zip("ABCDEFGH", (5, 32, 24, 8, 16, 16, 15, 10)):
        ws.column_dimensions[col].width = w
    wb.save(path); return path

def variazioni_word(d, path):
    doc = _doc("Variazioni prezzi — azioni in portafoglio",
               "Prezzo per azione in € · " + d["dal"] + " → " + d["al"] + " · var. include effetto cambio")
    p = doc.add_paragraph()
    p.add_run("Sintesi: ").bold = True
    best = (d["best"]["name"] + " " + pct(d["best"]["var"]*100)) if d["best"] else "—"
    worst = (d["worst"]["name"] + " " + pct(d["worst"]["var"]*100)) if d["worst"] else "—"
    p.add_run("media semplice " + pct(d["media"]*100) + " · media ponderata " + pct(d["media_w"]*100)
              + " · in rialzo " + str(d["up"]) + " / in ribasso " + str(d["down"])
              + " · migliore " + best + " · peggiore " + worst + ".")
    _table(doc, ["#", "Titolo", "Settore", "Div.", "Prezzo € " + d["dal"][:5], "Prezzo € " + d["al"][:5], "Var. %", "% NAV"],
           [[i, x["name"], x["settore"], x["ccy"], _pe(x["eur_dal"]), _pe(x["eur_al"]),
             (pct(x["var"]*100) if x["var"] is not None else "n.d."), pct(x["peso"]*100, 2, seg=False)]
            for i, x in enumerate(d["rows"], 1)])
    doc.save(path); return path

# -*- coding: utf-8 -*-
"""Report 4 - Sintesi Comitato (Word). Performance di periodo e YTD vs benchmark (SRE),
pesi per asset class AL vs DAL con i 3 fondi separati (WCTDD+VAL), esposizione azionaria
DELTA-ADJUSTED (dirette + delta azionario fondi - copertura derivati), settori, EM,
top/bottom performance, contributori/detrattori, operazioni del periodo (ORD). Dati da Oracle."""
import os
import data_layer as DL
import lookthrough as L

def _pct(x, dec=2, sign=False):
    if x is None: return "n.d."
    fmt = ("{:+." if sign else "{:.") + str(dec) + "f}"
    return fmt.format(x * 100).replace(".", ",") + "%"

def _eur(x):
    return "{:,.0f}".format(round(x or 0)).replace(",", ".") + " €"

def _num(x, dec=2):
    return ("{:,." + str(dec) + "f}").format(x or 0).replace(",", "§").replace(".", ",").replace("§", ".")

_DEV = {"EUR", "USD", "GBP", "CHF", "JPY", "NOK", "DKK", "SEK", "CAD", "Oro", "Altro (<1%)"}

def build_comitato(pf, repo_dir):
    meta = pf["meta"]
    var = L.build_variazioni(pf, repo_dir)
    mat = L.build_matrix(pf, repo_dir)
    tit = L.build_titoli(pf, repo_dir)
    ex = DL.comitato_extra(meta["schema"], meta["codcli"], meta.get("data_prec"), meta["data"])
    nav_al = ex.get("nav_al") or float(meta.get("nav") or 0) or 1.0
    nav_dal = ex.get("nav_dal") or nav_al
    # performance periodo + YTD
    ret = (ex["tcli_al"] / ex["tcli_dal"] - 1) if ex.get("tcli_dal") else None
    bmk = (ex["tbmk_al"] / ex["tbmk_dal"] - 1) if ex.get("tbmk_dal") else None
    dret = (ret - bmk) if (ret is not None and bmk is not None) else None
    rytd = (ex["tcli_al"] / ex["tcli_base"] - 1) if ex.get("tcli_base") else None
    bytd = (ex["tbmk_al"] / ex["tbmk_base"] - 1) if ex.get("tbmk_base") else None
    dytd = (rytd - bytd) if (rytd is not None and bytd is not None) else None
    # pesi asset class (fondi gia' separati da comitato_extra)
    comp = []
    for c in ex.get("comp", []):
        pa = c["val_al"] / nav_al; pd = c["val_dal"] / nav_dal
        comp.append({"macro": c["macro"], "pa": pa, "pd": pd, "d": pa - pd})
    comp.sort(key=lambda x: -x["pa"])
    # esposizione azionaria DELTA-ADJUSTED (motore Titoli)
    navt = tit.get("nav") or nav_al
    e_dir = tit.get("tot_diretto", 0.0) / navt
    e_fond = (tit.get("tot_fondi", 0.0) + tit.get("tot_indici", 0.0)) / navt
    e_hed = tit.get("tot_hedge", 0.0) / navt
    e_net = tit.get("tot_netto", 0.0) / navt
    em = sum(v for cc, v in mat["totale"].items() if cc not in _DEV)
    em_pct = em / mat["nav"] if mat.get("nav") else 0.0
    # settori (azioni dirette, vista Variazioni)
    eqrows = [r for r in var["rows"] if r["var"] is not None]
    sett = {}
    for r in eqrows:
        sett[r["settore"]] = sett.get(r["settore"], 0.0) + r["peso"]
    sett = sorted(sett.items(), key=lambda x: -x[1])
    # top/bottom performance e contributi
    byvar = sorted(eqrows, key=lambda r: -r["var"])
    top = byvar[:10]; bot = list(reversed(byvar[-10:]))
    contrib_class = sorted([{"macro": c["macro"], "contrib": (c["pa"] * nav_al - c["pd"] * nav_dal) / nav_dal}
                            for c in comp], key=lambda x: -x["contrib"])
    bycon = sorted(eqrows, key=lambda r: -(r["peso"] * r["var"]))
    ctop = bycon[:10]; cbot = list(reversed(bycon[-10:]))
    return {"dal": var["dal"], "al": var["al"], "bench": ex.get("bench", ""),
            "nav_al": nav_al, "nav_dal": nav_dal,
            "ret": ret, "bmk": bmk, "dret": dret, "rytd": rytd, "bytd": bytd, "dytd": dytd,
            "comp": comp, "e_dir": e_dir, "e_fond": e_fond, "e_hed": e_hed, "e_net": e_net, "em_pct": em_pct,
            "sett": sett, "top": top, "bot": bot, "ctop": ctop, "cbot": cbot,
            "trades": ex.get("trades", []), "best": var.get("best"), "worst": var.get("worst"),
            "contrib_class": contrib_class}

def comitato_word(d, path):
    from docx.shared import Pt, RGBColor
    doc = L._doc("Sintesi Comitato Investimenti — Linea Camperio",
                 "Periodo " + d["dal"] + " → " + d["al"] + " · benchmark " + (d["bench"] or "—")
                 + " · NAV " + _eur(d["nav_al"]) + " · USO INTERNO")
    def h(txt):
        p = doc.add_paragraph(); r = p.add_run(txt); r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x15, 0x1F, 0x6D)
    # 1. Performance
    h("1. Performance vs benchmark")
    L._table(doc, ["Orizzonte", "Portafoglio", "Benchmark " + (d["bench"] or ""), "Scostamento p.p."],
             [["Periodo (" + d["dal"] + " → " + d["al"] + ")", _pct(d["ret"], 2, True), _pct(d["bmk"], 2, True), _pct(d["dret"], 2, True)],
              ["Da inizio anno (YTD)", _pct(d["rytd"], 2, True), _pct(d["bytd"], 2, True), _pct(d["dytd"], 2, True)]])
    doc.add_paragraph("NAV da " + _eur(d["nav_dal"]) + " a " + _eur(d["nav_al"]) + ".")
    # 2. Pesi per asset class (fondi separati)
    h("2. Pesi di linea per asset class (vs " + d["dal"] + ")")
    L._table(doc, ["Asset class / fondo", "% NAV " + d["al"][:5], "% NAV " + d["dal"][:5], "Δ p.p."],
             [[c["macro"], _pct(c["pa"]), _pct(c["pd"]), _pct(c["d"], 2, True)] for c in d["comp"]]
             + [["TOTALE", "100,00%", "100,00%", ""]])
    # 3. Esposizione azionaria DELTA-ADJUSTED
    h("3. Esposizione azionaria delta-adjusted")
    L._table(doc, ["Componente", "% NAV"],
             [["Azioni dirette", _pct(d["e_dir"])],
              ["Azionario via fondi (delta: equity + indici)", _pct(d["e_fond"], 2, True)],
              ["Copertura derivati (hedge S&P short)", _pct(d["e_hed"], 2, True)],
              ["= Esposizione azionaria netta delta-adjusted", _pct(d["e_net"])],
              ["Di cui mercati emergenti (look-through)", _pct(d["em_pct"])]])
    # 4. Settori
    h("4. Peso per settore (azioni in portafoglio)")
    L._table(doc, ["Settore", "% NAV"], [[s, _pct(w)] for s, w in d["sett"]])
    # 5. Migliori/peggiori
    h("5. Migliori e peggiori per performance di periodo")
    L._table(doc, ["#", "Migliori 10", "Var. %", "Peggiori 10", "Var. %"],
             [[i + 1, d["top"][i]["name"], _pct(d["top"][i]["var"], 2, True),
               (d["bot"][i]["name"] if i < len(d["bot"]) else ""),
               (_pct(d["bot"][i]["var"], 2, True) if i < len(d["bot"]) else "")] for i in range(min(10, len(d["top"])))])
    # 6. Contributori/detrattori
    h("6. Contributori e detrattori di performance (peso × variazione, p.p.)")
    L._table(doc, ["#", "Top contributori", "Contrib. p.p.", "Top detrattori", "Contrib. p.p."],
             [[i + 1, d["ctop"][i]["name"], _pct(d["ctop"][i]["peso"] * d["ctop"][i]["var"], 3, True),
               (d["cbot"][i]["name"] if i < len(d["cbot"]) else ""),
               (_pct(d["cbot"][i]["peso"] * d["cbot"][i]["var"], 3, True) if i < len(d["cbot"]) else "")] for i in range(min(10, len(d["ctop"])))])
    doc.add_paragraph().add_run("Contributo per classe di attivo (p.p.) — include azioni, fondi, obbligazioni, oro, derivati").italic = True
    L._table(doc, ["Classe di attivo", "Contributo p.p."],
             [[c["macro"], _pct(c["contrib"], 3, True)] for c in d["contrib_class"]])
    # 7. Operazioni con segno
    h("7. Operazioni del periodo")
    if d["trades"]:
        L._table(doc, ["Data", "Titolo", "Operazione", "Quantità", "Prezzo", "Controvalore €"],
                 [[t["data"], t["nome"], t["verso"], _num(t["qty"], 0), _num(t["prezzo"], 2), _eur(t["ctv"])] for t in d["trades"]])
    else:
        doc.add_paragraph("Nessuna operazione di compravendita nel periodo (esclusi movimenti di cambio).")
    doc.save(path); return path

def comitato_html(d):
    def cc(x): return "1F7A4D" if (x or 0) >= 0 else "B3261E"
    head = ('<div class="rh"><div class="rt">Sintesi Comitato — Linea Camperio</div>'
            f'<div class="rs">Periodo {d["dal"]} &#8594; {d["al"]} · benchmark {d["bench"]} · NAV {_eur(d["nav_al"])} · export in Word</div></div>')
    kp = ('<div class="kp">'
          + f'<div class="ki"><div class="kl">Perf. periodo</div><div class="kv" style="color:#{cc(d["ret"])}">{_pct(d["ret"],2,True)}</div><div class="ks">bench {_pct(d["bmk"],2,True)} · Δ {_pct(d["dret"],2,True)} pp</div></div>'
          + f'<div class="ki"><div class="kl">Perf. YTD</div><div class="kv" style="color:#{cc(d["rytd"])}">{_pct(d["rytd"],2,True)}</div><div class="ks">bench {_pct(d["bytd"],2,True)} · Δ {_pct(d["dytd"],2,True)} pp</div></div>'
          + f'<div class="ki"><div class="kl">Azionario netto delta-adj</div><div class="kv">{_pct(d["e_net"])}</div><div class="ks">diretta {_pct(d["e_dir"])} · EM {_pct(d["em_pct"])}</div></div>'
          + f'<div class="ki"><div class="kl">Operazioni periodo</div><div class="kv">{len(d["trades"])}</div><div class="ks">con segno acq./vend.</div></div>'
          + '</div>')
    def tbl(headers, rows):
        th = "".join((f"<th class=r>{x}</th>" if i else f"<th>{x}</th>") for i, x in enumerate(headers))
        body = "".join("<tr>" + "".join((f"<td class=r>{c}</td>" if i else f"<td>{c}</td>") for i, c in enumerate(r)) + "</tr>" for r in rows)
        return f'<div style="overflow-x:auto"><table><thead><tr>{th}</tr></thead><tbody>{body}</tbody></table></div>'
    comp = tbl(["Asset class / fondo", "% " + d["al"][:5], "% " + d["dal"][:5], "Δ pp"],
               [[c["macro"], _pct(c["pa"]), _pct(c["pd"]), _pct(c["d"], 2, True)] for c in d["comp"]])
    eqt = tbl(["Componente", "% NAV"],
              [["Azioni dirette", _pct(d["e_dir"])],
               ["Azionario via fondi (delta)", _pct(d["e_fond"], 2, True)],
               ["Copertura derivati (hedge S&P)", _pct(d["e_hed"], 2, True)],
               ["= Azionario netto delta-adjusted", _pct(d["e_net"])],
               ["di cui Emerging Markets", _pct(d["em_pct"])]])
    sett = tbl(["Settore", "% NAV"], [[s, _pct(w)] for s, w in d["sett"]])
    perf = tbl(["#", "Migliori 10", "Var. %", "Peggiori 10", "Var. %"],
               [[i + 1, d["top"][i]["name"], _pct(d["top"][i]["var"], 2, True),
                 (d["bot"][i]["name"] if i < len(d["bot"]) else ""),
                 (_pct(d["bot"][i]["var"], 2, True) if i < len(d["bot"]) else "")] for i in range(min(10, len(d["top"])))])
    contr = tbl(["#", "Contributori", "p.p.", "Detrattori", "p.p."],
                [[i + 1, d["ctop"][i]["name"], _pct(d["ctop"][i]["peso"] * d["ctop"][i]["var"], 3, True),
                  (d["cbot"][i]["name"] if i < len(d["cbot"]) else ""),
                  (_pct(d["cbot"][i]["peso"] * d["cbot"][i]["var"], 3, True) if i < len(d["cbot"]) else "")] for i in range(min(10, len(d["ctop"])))])
    clazz = tbl(["Classe di attivo", "Contributo p.p."], [[c["macro"], _pct(c["contrib"], 3, True)] for c in d["contrib_class"]])
    trd = (tbl(["Data", "Titolo", "Operazione", "Quantità", "Prezzo", "Controvalore"],
               [[t["data"], t["nome"], t["verso"], _num(t["qty"], 0), _num(t["prezzo"], 2), _eur(t["ctv"])] for t in d["trades"]])
           if d["trades"] else '<div class="note">Nessuna operazione nel periodo (esclusi cambi).</div>')
    return (head + kp
            + "<h3>Pesi di linea per asset class (fondi separati)</h3>" + comp
            + "<h3>Esposizione azionaria delta-adjusted</h3>" + eqt
            + "<h3>Peso per settore (azioni)</h3>" + sett
            + "<h3>Migliori e peggiori per performance</h3>" + perf
            + "<h3>Contributori e detrattori (singoli titoli)</h3>" + contr
            + "<h3>Contributo per classe di attivo</h3>" + clazz
            + "<h3>Operazioni del periodo</h3>" + trd)
